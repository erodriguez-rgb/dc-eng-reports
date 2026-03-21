#!/usr/bin/env python3
"""
Generador de Informe Mensual DOCX — DC Engineering Group
Portada moderna, colores DC Engineering, logos, TOC con números de página.
"""

import io, os, requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── PALETA DC ENGINEERING ─────────────────────────────────────────────────────
GREEN_DARK   = "2E7D32"   # verde oscuro encabezados
GREEN_MID    = "4CAF50"   # verde medio
GREEN_LIGHT  = "C8E6C9"   # verde claro fondos de tabla
TEAL         = "00796B"   # teal acento
WHITE        = "FFFFFF"
DARK_TEXT    = "1A1A1A"
GRAY_LIGHT   = "F5F5F5"

# RGB tuples
RGB_GREEN_DARK  = RGBColor(0x2E, 0x7D, 0x32)
RGB_GREEN_MID   = RGBColor(0x4C, 0xAF, 0x50)
RGB_TEAL        = RGBColor(0x00, 0x79, 0x6B)
RGB_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
RGB_DARK        = RGBColor(0x1A, 0x1A, 0x1A)
RGB_GRAY        = RGBColor(0x75, 0x75, 0x75)

# ── HELPERS ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="000000", size="4", val="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), val)
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def no_borders_cell(cell):
    set_cell_borders(cell, val="none", color="FFFFFF", size="0")

def set_para_spacing(p, before=0, after=60):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)

def add_run(p, text, bold=False, italic=False, size=10, color=None, font="Arial"):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_para(doc, text="", bold=False, italic=False, size=10,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             before=0, after=60):
    p = doc.add_paragraph()
    p.alignment = alignment
    set_para_spacing(p, before, after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p

def section_heading(doc, roman, title):
    """Encabezado de sección con barra verde izquierda simulada via tabla."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.allow_autofit = False
    # Barra de color
    bar = t.cell(0, 0)
    bar.width = Inches(0.15)
    set_cell_bg(bar, GREEN_DARK)
    no_borders_cell(bar)
    bar.paragraphs[0].text = ""
    # Texto
    txt = t.cell(0, 1)
    no_borders_cell(txt)
    set_cell_bg(txt, GRAY_LIGHT)
    p = txt.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, 60, 60)
    add_run(p, f"  {roman}.  ", bold=True, size=11, color=RGB_GREEN_DARK)
    add_run(p, title.upper(), bold=True, size=11, color=RGB_DARK)
    # Espacio después
    doc.add_paragraph()

def table_header_row(table, headers, bg=GREEN_DARK, text_color=RGB_WHITE):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i >= len(row.cells): break
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, 60, 60)
        add_run(p, h, bold=True, size=9, color=text_color)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, color=WHITE, size="4")

def data_row(table, values, alt=False):
    row = table.add_row()
    bg = "F1F8E9" if alt else WHITE
    for i, v in enumerate(values):
        if i >= len(row.cells): break
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_para_spacing(p, 40, 40)
        add_run(p, str(v) if v is not None else "—", size=9)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, color="CCCCCC", size="4")
    return row

def fetch_image(url):
    """Descarga imagen desde URL, retorna BytesIO o None."""
    if not url:
        return None
    try:
        r = requests.get(url, timeout=10)
        if r.status_code == 200:
            return io.BytesIO(r.content)
    except Exception:
        pass
    return None

def add_page_number(doc):
    """Agrega números de página al pie centrado."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    # Número de página via campo
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    run2._r.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld2)
    # Estilo
    for r in [run, run2, run3]:
        r.font.size = Pt(9)
        r.font.color.rgb = RGB_GRAY

# ── PORTADA MODERNA ───────────────────────────────────────────────────────────

def build_cover_page(doc, data, dc_img, ag_img):
    """
    Portada con:
    - Banda superior verde oscuro con logos
    - Nombre del proyecto en verde grande
    - Línea decorativa
    - Info del informe
    - Banda inferior con info DC Engineering
    """
    section = doc.sections[0]

    # ── Banda superior: tabla 1 fila, fondo verde oscuro ──
    t_top = doc.add_table(rows=1, cols=3)
    t_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_top.allow_autofit = False
    # Anchos: logo agencia | espacio | logo DC
    col_w = [Inches(2.0), Inches(2.5), Inches(2.0)]
    for i, w in enumerate(col_w):
        for cell in t_top.columns[i].cells:
            cell.width = w

    for cell in t_top.rows[0].cells:
        set_cell_bg(cell, GREEN_DARK)
        no_borders_cell(cell)

    # Logo agencia (izquierda)
    cell_ag = t_top.cell(0, 0)
    if ag_img:
        try:
            p = cell_ag.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(ag_img, width=Inches(1.4))
        except:
            p = cell_ag.paragraphs[0]
            add_run(p, data.get('agencia', ''), bold=True, size=10, color=RGB_WHITE)
    else:
        p = cell_ag.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, 120, 40)
        add_run(p, data.get('agencia', ''), bold=True, size=10, color=RGB_WHITE)

    # Espacio central
    t_top.cell(0, 1).paragraphs[0].text = ""

    # Logo DC Engineering (derecha)
    cell_dc = t_top.cell(0, 2)
    if dc_img:
        try:
            p = cell_dc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run()
            run.add_picture(dc_img, width=Inches(1.4))
        except:
            p = cell_dc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, "DC ENGINEERING GROUP", bold=True, size=9, color=RGB_WHITE)
    else:
        p = cell_dc.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p, 120, 40)
        add_run(p, "DC ENGINEERING GROUP", bold=True, size=9, color=RGB_WHITE)

    # Espacio
    sp = doc.add_paragraph()
    set_para_spacing(sp._p.get_or_add_pPr() and sp or sp, 0, 0)
    sp.paragraph_format.space_after = Pt(18)

    # ── Nombre del proyecto ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, 120, 60)
    add_run(p, data.get('nombre_proyecto', '').upper(),
            bold=True, size=20, color=RGB_GREEN_DARK)

    # Línea decorativa verde
    t_line = doc.add_table(rows=1, cols=1)
    t_line.allow_autofit = False
    cell_line = t_line.cell(0, 0)
    cell_line.width = Inches(6.5)
    set_cell_bg(cell_line, GREEN_MID)
    no_borders_cell(cell_line)
    p2 = cell_line.paragraphs[0]
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Info del informe ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, 40, 40)
    num = data.get('numero_informe', 1)
    periodo = data.get('periodo_texto', '')
    add_run(p, f"Informe Mensual #{num}:  ", bold=True, size=11, color=RGB_TEAL)
    add_run(p, periodo, bold=False, size=11, color=RGB_DARK)

    p3 = doc.add_paragraph()
    set_para_spacing(p3, 20, 20)
    add_run(p3, f"Contrato: {data.get('numero_contrato', 'N/A')}   |   Contratista: {data.get('contratista', 'N/A')}",
            size=10, color=RGB_GRAY)

    # Espacio grande antes de banda inferior
    for _ in range(8):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)

    # ── Banda inferior: fondo verde oscuro con info DC ──
    t_bot = doc.add_table(rows=1, cols=2)
    t_bot.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_bot.allow_autofit = False
    t_bot.cell(0, 0).width = Inches(3.5)
    t_bot.cell(0, 1).width = Inches(3.0)

    for cell in t_bot.rows[0].cells:
        set_cell_bg(cell, GREEN_DARK)
        no_borders_cell(cell)

    p_left = t_bot.cell(0, 0).paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_left, 100, 40)
    add_run(p_left, "DC ENGINEERING GROUP, PSC\n", bold=True, size=11, color=RGB_WHITE)
    add_run(p_left, "First Federal Savings Building Suite #510\nPonce de León Ave. #1519, SJ, PR, 00909", size=9, color=RGB_WHITE)

    p_right = t_bot.cell(0, 1).paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p_right, 100, 40)
    add_run(p_right, "Tel: 787-477-1789\ndcordova@dc-eng.com", size=9, color=RGB_WHITE)

    doc.add_page_break()


# ── TABLA DE CONTENIDO CON NÚMEROS DE PÁGINA ─────────────────────────────────

TOC_ITEMS = [
    ("I.",     "Descripción del Proyecto",                    3),
    ("II.",    "Información General del Proyecto",            3),
    ("III.",   "Resumen de Trabajos Realizados",              3),
    ("IV.",    "Análisis de Costos y Tiempo",                 4),
    ("V.",     "Fianzas y Seguros",                           4),
    ("VI.",    "Resumen de Facturación",                      5),
    ("VII.",   "Request for Information",                     5),
    ("VIII.",  "Permisos y Endosos",                          5),
    ("IX.",    "Submittal",                                   6),
    ("X.",     "Anejo A: Comunicación Escrita Enviada/Recibida", 7),
    ("XI.",    "Anejo B: Submittals",                        39),
    ("XII.",   "Anejo C: Informes Diarios de Inspección",    51),
]

def build_toc(doc, ag_img):
    """Tabla de contenido estilizada con número de página."""
    # Mini header con logo agencia
    if ag_img:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(ag_img, width=Inches(0.8))
        except:
            pass

    # Encabezado TOC
    t_hdr = doc.add_table(rows=1, cols=1)
    t_hdr.allow_autofit = False
    cell = t_hdr.cell(0, 0)
    set_cell_bg(cell, GREEN_DARK)
    no_borders_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, 80, 80)
    add_run(p, "  TABLE OF CONTENTS", bold=True, size=12, color=RGB_WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Filas del TOC con dots
    for num, titulo, pagina in TOC_ITEMS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, 30, 30)
        # Tab stops para alinear número de página
        add_run(p, f"{num} ", bold=True, size=10, color=RGB_GREEN_DARK)
        # Calcular dots
        label = f"{titulo}"
        dots = "." * max(3, 60 - len(titulo) - len(num))
        add_run(p, label, size=10, color=RGB_DARK)
        add_run(p, dots, size=10, color=RGB_GRAY)
        add_run(p, str(pagina), bold=True, size=10, color=RGB_GREEN_DARK)

    doc.add_page_break()


# ── FUNCIÓN PRINCIPAL ─────────────────────────────────────────────────────────

def generar_informe_mensual(data):
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Números de página en footer
    add_page_number(doc)

    # Descargar logos
    dc_img = fetch_image("https://media.base44.com/images/public/69b83cf2a12fa0a64f315f09/c53792645_DCENGINEERING.png")
    ag_img = fetch_image(data.get('agencia_logo_url', ''))

    # ── PORTADA ──
    build_cover_page(doc, data, dc_img, ag_img)

    # Reiniciar logo para TOC
    if ag_img:
        ag_img.seek(0)
    if dc_img:
        dc_img.seek(0)
    ag_img2 = fetch_image(data.get('agencia_logo_url', ''))

    # ── TABLA DE CONTENIDO ──
    build_toc(doc, ag_img2)

    contratista = data.get('contratista', '')

    # ── I. DESCRIPCIÓN DEL PROYECTO ──────────────────────────────────────────
    section_heading(doc, "I", "Descripción del Proyecto")
    nombre_proy = data.get('nombre_proyecto', '')
    add_para(doc, f"El proyecto propuesto consiste en la {nombre_proy}.", size=10, before=0, after=80)

    # ── II. INFORMACIÓN GENERAL DEL PROYECTO ─────────────────────────────────
    section_heading(doc, "II", "Información General del Proyecto")
    add_para(doc, "Partes Involucradas en el Proyecto de Referencia:", bold=True, size=10, after=40)

    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    table_header_row(t, ['Participantes', 'Descripción'])

    inspector = data.get('inspector', 'Ing. Francisco Cordova, PE')
    partes = [
        ('Dueño', f"{data.get('agencia', 'Municipio Autónomo de Caguas')}\nAgrim. Yamil Sanabria Roman\nDirector Interino — Departamento de Obras Públicas\nEmail: Yamil.Sanabria@caguas.gov.pr"),
        ('Inspección', f"DC ENGINEERING GROUP, PSC\nFirst Federal Savings Building Suite #510, Ponce de León Ave. #1519, SJ, PR, 00909\nProject Inspector: {inspector} — Email: fcordova@dc-eng.com\nTel: 787-318-0746"),
    ]
    if data.get('disenador'):
        partes.append(('Diseño', data['disenador']))
    partes.append(('Contratista General', f"{contratista}\n{data.get('contratista_contacto', '')}".strip()))

    for i, (rol, desc) in enumerate(partes):
        row = data_row(t, [rol, desc], alt=(i % 2 == 0))
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── III. RESUMEN DE TRABAJOS REALIZADOS ───────────────────────────────────
    section_heading(doc, "III", "Resumen de Trabajos Realizados")

    add_para(doc, "Comunicación Escrita", bold=True, size=10, color=RGB_TEAL, after=30)
    comunicaciones = data.get('comunicaciones', [])
    if comunicaciones:
        add_para(doc, "Durante el periodo que comprende este informe se generaron las siguientes comunicaciones escritas:", size=10, after=30)
        for com in comunicaciones:
            p = doc.add_paragraph(style='List Bullet')
            set_para_spacing(p, 20, 20)
            add_run(p, com, size=10)
    else:
        add_para(doc, "Durante el periodo que comprende este informe no se recibieron y/o enviaron comunicaciones escritas.", size=10, after=30)
    add_para(doc, "Para más detalles ver Anejo A – Comunicaciones escritas", italic=True, size=9, color=RGB_GRAY, after=60)

    add_para(doc, "Tareas De Campo:", bold=True, size=10, color=RGB_TEAL, after=30)
    resumen_act = data.get('resumen_actividades', '')
    for linea in resumen_act.split('\n'):
        linea = linea.strip().lstrip('•').strip()
        if linea:
            p = doc.add_paragraph(style='List Bullet')
            set_para_spacing(p, 20, 20)
            add_run(p, linea, size=10)
    add_para(doc, "Para más detalles ver Anejo C – Informes Semanales y Diarios", italic=True, size=9, color=RGB_GRAY, after=60)

    # ── IV. ANÁLISIS DE COSTOS Y TIEMPO ──────────────────────────────────────
    section_heading(doc, "IV", "Análisis de Costos y Tiempo")
    add_para(doc, f"A continuación, se detalla el análisis de costos y tiempo del contratista {contratista}.", size=10, after=40)
    add_para(doc, contratista, bold=True, size=10, color=RGB_TEAL, after=30)
    add_para(doc, f"Número de Contrato: {data.get('numero_contrato', 'N/A')}", size=10, after=40)

    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.columns[0].width = Inches(1.5)
    t.columns[1].width = Inches(5.0)
    table_header_row(t, ['Concepto', 'Detalle'])
    data_row(t, ['COST', f"PO {data.get('po_number', 'N/A')} — {data.get('costo_contrato', 'N/A')}"], alt=True)
    data_row(t, ['TIME', f"Duración del contrato: {data.get('dias_contrato', 'N/A')}"])
    doc.add_paragraph()

    # ── V. FIANZAS Y SEGUROS ──────────────────────────────────────────────────
    section_heading(doc, "V", "Fianzas y Seguros")
    add_para(doc, f"Las tablas a continuación resumen los seguros del contratista {contratista}.", size=10, after=40)
    add_para(doc, contratista, bold=True, size=10, color=RGB_TEAL, after=30)

    cols_seg = ['Tipo De Póliza', 'Número', 'Fecha Emisión', 'Fecha Venc.', 'Límite']
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    widths_seg = [1.5, 1.2, 1.0, 1.0, 1.8]
    for i, w in enumerate(widths_seg):
        t.columns[i].width = Inches(w)
    table_header_row(t, cols_seg)
    seguros = data.get('seguros', [])
    if seguros:
        for i, seg in enumerate(seguros):
            data_row(t, [seg.get('tipo','—'), seg.get('numero','—'), seg.get('fecha_emision','—'), seg.get('fecha_vencimiento','—'), seg.get('limite','—')], alt=(i%2==0))
    else:
        data_row(t, ['—','—','—','—','—'])
    doc.add_paragraph()

    # ── VI. RESUMEN DE FACTURACIÓN ────────────────────────────────────────────
    section_heading(doc, "VI", "Resumen de Facturación")
    add_para(doc, f"Las tablas a continuación resumen las facturas del contratista {contratista}.", size=10, after=40)
    add_para(doc, contratista, bold=True, size=10, color=RGB_TEAL, after=30)

    cols_fac = ['No.', 'Contratista', 'Cantidad', 'Porciento', 'Fecha', 'Status']
    t = doc.add_table(rows=1, cols=6)
    t.style = 'Table Grid'
    widths_fac = [0.4, 2.0, 0.9, 0.8, 1.1, 1.3]
    for i, w in enumerate(widths_fac):
        t.columns[i].width = Inches(w)
    table_header_row(t, cols_fac)
    certificaciones = data.get('certificaciones', [])
    if certificaciones:
        for i, cert in enumerate(certificaciones):
            data_row(t, [cert.get('numero','—'), cert.get('contratista', contratista), cert.get('cantidad','—'), cert.get('porciento','—'), cert.get('fecha','—'), cert.get('status','Pendiente')], alt=(i%2==0))
    else:
        data_row(t, ['1', contratista, '—', '—', '—', '—'])
    doc.add_paragraph()

    # ── VII. REQUEST FOR INFORMATION ──────────────────────────────────────────
    section_heading(doc, "VII", "Request for Information")
    rfis = data.get('rfis', [])
    if rfis:
        add_para(doc, "Durante el periodo se recibieron los siguientes RFIs:", size=10, after=40)
    else:
        add_para(doc, "Durante el periodo que comprende este informe, no se recibieron RFIs.", size=10, after=40)
    add_para(doc, contratista, bold=True, size=10, color=RGB_TEAL, after=30)

    cols_rfi = ['No.', 'RFI No.', 'Description', 'Date Sent']
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    t.columns[0].width = Inches(0.4)
    t.columns[1].width = Inches(0.8)
    t.columns[2].width = Inches(3.8)
    t.columns[3].width = Inches(1.5)
    table_header_row(t, cols_rfi)
    if rfis:
        for i, rfi in enumerate(rfis):
            data_row(t, [rfi.get('no','—'), rfi.get('rfiNo','—'), rfi.get('descripcion','—'), rfi.get('fecha','—')], alt=(i%2==0))
    else:
        data_row(t, ['1', '—', '—', '—'])
    doc.add_paragraph()

    # ── VIII. PERMISOS Y ENDOSOS ──────────────────────────────────────────────
    section_heading(doc, "VIII", "Permisos y Endosos")
    add_para(doc, "Durante el periodo que comprende este informe no se generaron permisos.", size=10, after=40)

    cols_per = ['Descripción', 'Agencia', 'Estatus', 'Número de Caso']
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    t.columns[0].width = Inches(2.5)
    t.columns[1].width = Inches(1.5)
    t.columns[2].width = Inches(1.0)
    t.columns[3].width = Inches(1.5)
    table_header_row(t, cols_per)
    data_row(t, ['—', '—', '—', '—'])
    doc.add_paragraph()

    # ── IX. SUBMITTAL ─────────────────────────────────────────────────────────
    section_heading(doc, "IX", "Submittal")
    add_para(doc, "La tabla a continuación contiene todos los submittals del proyecto.", size=10, after=40)
    add_para(doc, contratista, bold=True, size=10, color=RGB_TEAL, after=30)

    cols_sub = ['No.', 'Sub. No.', 'Title', 'Status', 'Received', 'Sent Designer', 'Returned', 'Sent Contractor']
    t = doc.add_table(rows=1, cols=8)
    t.style = 'Table Grid'
    widths_sub = [0.35, 0.55, 1.8, 0.65, 0.75, 0.75, 0.75, 0.75]
    for i, w in enumerate(widths_sub):
        t.columns[i].width = Inches(w)
    table_header_row(t, cols_sub)
    submittals = data.get('submittals_acumulativos', [])
    if submittals:
        for i, s in enumerate(submittals):
            data_row(t, [i+1, s.get('no',i+1), s.get('titulo','—'), s.get('status','—'), s.get('fechaRecibido','—'), s.get('fechaDiseñador','—'), s.get('fechaRetornado','—'), s.get('fechaContratista','—')], alt=(i%2==0))
    else:
        data_row(t, ['1', '—', '—', '—', '—', '—', '—', '—'])
    doc.add_paragraph()

    # ── X. ANEJO A ────────────────────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc, "X", "Anejo A: Comunicación Escrita Enviada/Recibida")
    if comunicaciones:
        for com in comunicaciones:
            p = doc.add_paragraph(style='List Number')
            set_para_spacing(p, 20, 20)
            add_run(p, com, size=10)
    else:
        add_para(doc, "Durante el periodo que comprende este informe no se recibieron comunicaciones escritas.", size=10)

    # ── XI. ANEJO B ───────────────────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc, "XI", "Anejo B: Submittals")
    submittals_periodo = data.get('submittals_periodo', [])
    if submittals_periodo:
        for s in submittals_periodo:
            p = doc.add_paragraph(style='List Number')
            set_para_spacing(p, 20, 20)
            add_run(p, s, size=10)
    else:
        add_para(doc, "Durante el periodo que comprende este informe no se recibieron Submittals.", size=10)

    # ── XII. ANEJO C ──────────────────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc, "XII", "Anejo C: Informes Diarios de Inspección")
    add_para(doc, f"Se adjuntan los informes diarios de inspección correspondientes al período {data.get('periodo_texto', '')}.", size=10, after=60)

    actividades = data.get('actividades', [])
    if actividades:
        cols_act = ['Fecha', 'Descripción de Actividades', 'Observaciones']
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        t.columns[0].width = Inches(1.0)
        t.columns[1].width = Inches(3.2)
        t.columns[2].width = Inches(2.3)
        table_header_row(t, cols_act)
        for i, act in enumerate(actividades):
            data_row(t, [act.get('fecha','—'), act.get('descripcion','—'), act.get('observacion','—')], alt=(i%2==0))
    else:
        add_para(doc, "(Adjuntar PDF de informes diarios del período)", italic=True, size=10, color=RGB_GRAY)

    # Guardar
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
