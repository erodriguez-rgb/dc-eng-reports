#!/usr/bin/env python3
"""
Generador de Informe Diario DOCX
"""

import sys, json, io, os, urllib.request
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm, Emu
from docx.oxml import OxmlElement
from lxml import etree

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'template_informe.docx')

def fetch_image(url):
    """Descarga imagen y retorna bytes"""
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=15) as r:
            return r.read()
    except Exception as e:
        print(f"[WARN] No se pudo descargar {url}: {e}", file=sys.stderr)
        return None

def add_inline_image(para, img_bytes, width_emu, height_emu):
    """Crea un inline drawing element con la imagen"""
    image_stream = io.BytesIO(img_bytes)
    pic_rId = para.part.relate_to(
        image_stream,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
        is_external=False
    )

    import random
    doc_pr_id = random.randint(200, 9999)

    inline_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:rPr><w:noProof/></w:rPr>
      <w:drawing>
        <wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   distT="0" distB="0" distL="0" distR="0">
          <wp:extent cx="{width_emu}" cy="{height_emu}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:docPr id="{doc_pr_id}" name="Photo{doc_pr_id}"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:nvPicPr>
                  <pic:cNvPr id="0" name="photo"/>
                  <pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{pic_rId}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                  <a:stretch><a:fillRect/></a:stretch>
                </pic:blipFill>
                <pic:spPr bwMode="auto">
                  <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                  <a:noFill/>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:inline>
      </w:drawing>
    </w:r>'''

    run_elem = etree.fromstring(inline_xml)
    para._p.append(run_elem)
    return pic_rId

def generar_informe(data):
    """Genera el DOCX con los datos del informe"""
    doc = Document(TEMPLATE_PATH)

    # === PAGE 1: Text fields ===
    para5 = doc.paragraphs[5]
    value_runs = [r for r in para5.runs if not r.bold]
    if value_runs:
        value_runs[0].text = ': ' + data.get('nombre_proyecto', '')
        for r in value_runs[1:]: r.text = ''

    para6 = doc.paragraphs[6]
    for run in para6.runs:
        if not run.bold:
            run.text = ' ' + data.get('numero_contrato', '')
            break

    para7 = doc.paragraphs[7]
    value_runs = [r for r in para7.runs if not r.bold]
    if value_runs:
        value_runs[0].text = ' ' + data.get('contratista', '')
        for r in value_runs[1:]: r.text = ''

    para8 = doc.paragraphs[8]
    value_runs = [r for r in para8.runs if not r.bold]
    if value_runs:
        value_runs[0].text = ' ' + data.get('localizacion', '')
        for r in value_runs[1:]: r.text = ''

    para14 = doc.paragraphs[14]
    fecha_found = False
    cleared_after_fecha = False
    for run in para14.runs:
        if 'Fecha' in run.text:
            fecha_found = True
            continue
        if fecha_found and run.text.strip() in [':', ' ', '']:
            continue
        if fecha_found and not cleared_after_fecha:
            run.text = ' ' + data.get('fecha', '')
            cleared_after_fecha = True
            fecha_found = False
        elif cleared_after_fecha:
            run.text = ''

    # === TABLE 0: Weather ===
    weather_table = doc.tables[0]
    # Soportar cielo_am/cielo_pm separados O cielo combinado
    cielo_am = data.get('cielo_am', data.get('cielo', 'Soleado'))
    cielo_pm = data.get('cielo_pm', data.get('cielo', 'Soleado'))

    # Cell[1] = AM, Cell[2] = PM (tabla 5x3 del template real)
    for cell_idx, cielo_val in [(1, cielo_am), (2, cielo_pm)]:
        cell = weather_table.rows[1].cells[cell_idx]
        for para in cell.paragraphs:
            for run in para.runs: run.text = ''
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = cielo_val
        else:
            cell.paragraphs[0].add_run(cielo_val)

    lluvia_cell = weather_table.rows[2].cells[1]
    lluvia_val = data.get('lluvia_texto', 'Sí' if data.get('lluvia_ocurrio') else 'No')
    for para in lluvia_cell.paragraphs:
        for run in para.runs: run.text = ''
    if lluvia_cell.paragraphs[0].runs:
        lluvia_cell.paragraphs[0].runs[0].text = lluvia_val
    else:
        lluvia_cell.paragraphs[0].add_run(lluvia_val)

    pluv_cell = weather_table.rows[2].cells[2]
    pluv_runs = list(pluv_cell.paragraphs[0].runs)
    if len(pluv_runs) >= 2:
        pluv_runs[-1].text = ' ' + data.get('pluviometro', '0.0"')
    elif pluv_runs:
        pluv_runs[0].text = 'Pluviómetro: ' + data.get('pluviometro', '0.0"')

    dur_cell = weather_table.rows[3].cells[0]
    for para in dur_cell.paragraphs:
        for run in para.runs: run.text = ''
    if dur_cell.paragraphs[0].runs:
        dur_cell.paragraphs[0].runs[0].text = 'Duración: ' + data.get('lluvia_duracion', 'N/A')
    else:
        dur_cell.paragraphs[0].add_run('Duración: ' + data.get('lluvia_duracion', 'N/A'))

    # Soportar impacto_am/impacto_pm separados O impacto combinado
    impacto_am = data.get('impacto_am', data.get('lluvia_impacto', 'No tuvo impacto.'))
    impacto_pm = data.get('impacto_pm', data.get('lluvia_impacto', 'No tuvo impacto.'))
    for cell_idx, impacto_val in [(1, impacto_am), (2, impacto_pm)]:
        cell = weather_table.rows[4].cells[cell_idx]
        for para in cell.paragraphs:
            for run in para.runs: run.text = ''
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = impacto_val
        else:
            cell.paragraphs[0].add_run(impacto_val)

    # === TABLE 1: Activities ===
    act_table = doc.tables[1]

    def set_cell_text(row_idx, text):
        cell = act_table.rows[row_idx].cells[0]
        for para in cell.paragraphs:
            for run in para.runs: run.text = ''
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = text
        else:
            cell.paragraphs[0].add_run(text)

    set_cell_text(1, data.get('actividades', ''))
    set_cell_text(3, data.get('mano_obra', ''))
    set_cell_text(5, data.get('visitantes', ''))
    set_cell_text(7, data.get('observaciones', 'Sin comentarios.'))

    # === PAGE 2: Photos ===
    fotos = data.get('fotos', [])
    photo_para_indices = [23, 25, 35, 37]

    for i, para_idx in enumerate(photo_para_indices):
        if para_idx >= len(doc.paragraphs):
            break
        para = doc.paragraphs[para_idx]

        for drawing in para._p.findall('.//' + qn('w:drawing')):
            drawing.getparent().remove(drawing)
        for run in para.runs:
            run.text = ''

        photo_batch = fotos[i*2:i*2+2]

        for j, foto_url in enumerate(photo_batch):
            img_bytes = fetch_image(foto_url)
            if img_bytes:
                width_emu = int(3.4 * 914400)
                height_emu = int(2.5 * 914400)
                try:
                    add_inline_image(para, img_bytes, width_emu, height_emu)
                    print(f"[OK] Foto {i*2+j+1} agregada", file=sys.stderr)
                except Exception as e:
                    print(f"[ERROR] Foto {i*2+j+1}: {e}", file=sys.stderr)

    output_path = '/tmp/informe_output.docx'
    doc.save(output_path)
    return output_path
